"""
Document Text Extractor
-----------------------

Extracts text from:

- PDF files
- DOCX files
- Images (JPG, JPEG, PNG)

Features:

- Direct text extraction from normal PDFs
- OCR fallback for scanned/image-based PDFs
- OCR extraction from certificate images
- DOCX text extraction

Used by:
- upload_routes.py
- skill_extractor.py
"""

import os

import fitz
from docx import Document
from PIL import Image
import pytesseract


# ============================================================
# TESSERACT OCR CONFIGURATION
# ============================================================

TESSERACT_PATH = (
    r"C:\Program Files\Tesseract-OCR\tesseract.exe"
)


if os.path.exists(
    TESSERACT_PATH
):

    pytesseract.pytesseract.tesseract_cmd = (
        TESSERACT_PATH
    )

    print(
        "Tesseract OCR configured successfully."
    )

else:

    print(
        "WARNING: Tesseract OCR executable not found."
    )

    print(
        "Expected location:"
    )

    print(
        TESSERACT_PATH
    )


# ============================================================
# SUPPORTED IMAGE TYPES
# ============================================================

SUPPORTED_IMAGE_EXTENSIONS = {

    "jpg",

    "jpeg",

    "png"

}


# ============================================================
# CHECK TESSERACT
# ============================================================

def check_tesseract():

    """
    Check whether Tesseract OCR
    is available and working.
    """

    try:

        version = (
            pytesseract.get_tesseract_version()
        )

        print(
            f"Tesseract OCR Version: {version}"
        )

        return True

    except Exception as error:

        print(
            "Tesseract OCR is not available."
        )

        print(
            error
        )

        return False


# ============================================================
# EXTRACT TEXT FROM IMAGE
# ============================================================

def extract_text_from_image(file_path):

    """
    Extract text from an image using OCR.

    Supported:

    - JPG
    - JPEG
    - PNG
    """

    if not check_tesseract():

        return ""

    try:

        print(
            "Running OCR on image..."
        )


        image = Image.open(
            file_path
        )


        # Convert image to RGB
        #
        # This prevents some OCR issues
        # with PNG transparency.

        if image.mode != "RGB":

            image = image.convert(
                "RGB"
            )


        text = pytesseract.image_to_string(

            image,

            config="--psm 6"

        )


        return text.strip()


    except Exception as error:

        print(
            f"OCR image extraction error: {error}"
        )

        return ""


# ============================================================
# EXTRACT TEXT FROM NORMAL PDF
# ============================================================

def extract_text_from_pdf(file_path):

    """
    Extract text directly from PDF.

    If no selectable text exists,
    OCR fallback is used.
    """

    extracted_text = ""

    document = None


    try:

        print(
            "Trying direct PDF text extraction..."
        )


        document = fitz.open(
            file_path
        )


        for page_number, page in enumerate(
            document,
            start=1
        ):

            page_text = page.get_text()

            if page_text:

                extracted_text += (
                    page_text
                    + "\n"
                )


            print(
                f"PDF page {page_number} processed."
            )


    except Exception as error:

        print(
            f"PDF text extraction error: {error}"
        )

        return ""


    finally:

        if document:

            document.close()


    # --------------------------------------------------------
    # DIRECT TEXT FOUND
    # --------------------------------------------------------

    if extracted_text.strip():

        print(
            "Direct PDF text extraction successful."
        )

        return extracted_text.strip()


    # --------------------------------------------------------
    # OCR FALLBACK
    # --------------------------------------------------------

    print(
        "No selectable PDF text found."
    )

    print(
        "Trying OCR on scanned PDF..."
    )


    return extract_text_from_scanned_pdf(
        file_path
    )


# ============================================================
# OCR FOR SCANNED PDF
# ============================================================

def extract_text_from_scanned_pdf(file_path):

    """
    Convert scanned PDF pages into images
    and extract text using OCR.
    """

    if not check_tesseract():

        return ""


    extracted_text = ""

    document = None


    try:

        document = fitz.open(
            file_path
        )


        total_pages = len(
            document
        )


        print(
            f"Running OCR on {total_pages} PDF page(s)..."
        )


        for page_number in range(
            total_pages
        ):

            page = document.load_page(
                page_number
            )


            # ------------------------------------------------
            # HIGHER RESOLUTION
            #
            # Better OCR accuracy
            # ------------------------------------------------

            matrix = fitz.Matrix(
                2,
                2
            )


            pixmap = page.get_pixmap(

                matrix=matrix,

                alpha=False

            )


            image = Image.frombytes(

                "RGB",

                (
                    pixmap.width,
                    pixmap.height
                ),

                pixmap.samples

            )


            page_text = (
                pytesseract.image_to_string(

                    image,

                    config="--psm 6"

                )
            )


            if page_text:

                extracted_text += (

                    page_text

                    + "\n"

                )


            print(

                f"OCR completed for page "
                f"{page_number + 1}/{total_pages}"

            )


        return extracted_text.strip()


    except Exception as error:

        print(
            f"Scanned PDF OCR error: {error}"
        )

        return ""


    finally:

        if document:

            document.close()


# ============================================================
# EXTRACT TEXT FROM DOCX
# ============================================================

def extract_text_from_docx(file_path):

    """
    Extract text from DOCX files.
    """

    try:

        print(
            "Extracting text from DOCX..."
        )


        document = Document(
            file_path
        )


        paragraphs = []


        # ----------------------------------------------------
        # NORMAL PARAGRAPHS
        # ----------------------------------------------------

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:

                paragraphs.append(
                    text
                )


        # ----------------------------------------------------
        # TABLE CONTENT
        # ----------------------------------------------------

        for table in document.tables:

            for row in table.rows:

                for cell in row.cells:

                    text = cell.text.strip()

                    if text:

                        paragraphs.append(
                            text
                        )


        extracted_text = "\n".join(
            paragraphs
        )


        print(
            "DOCX extraction completed."
        )


        return extracted_text.strip()


    except Exception as error:

        print(
            f"DOCX extraction error: {error}"
        )

        return ""


# ============================================================
# MAIN FILE EXTRACTOR
# ============================================================

def extract_text_from_file(file_path):

    """
    Automatically detect file type
    and extract text.

    Supported:

    - PDF
    - DOCX
    - JPG
    - JPEG
    - PNG

    Returns:

    str
    """

    # --------------------------------------------------------
    # VALIDATE PATH
    # --------------------------------------------------------

    if not file_path:

        print(
            "No file path provided."
        )

        return ""


    if not os.path.exists(
        file_path
    ):

        print(
            f"File does not exist: {file_path}"
        )

        return ""


    # --------------------------------------------------------
    # GET EXTENSION
    # --------------------------------------------------------

    extension = (

        os.path.splitext(
            file_path
        )[1]

        .lower()

        .replace(
            ".",
            ""
        )

    )


    print(
        "\n========================================"
    )

    print(
        f"EXTRACTING TEXT FROM: "
        f"{extension.upper()}"
    )

    print(
        f"FILE: "
        f"{os.path.basename(file_path)}"
    )

    print(
        "========================================"
    )


    # --------------------------------------------------------
    # PDF
    # --------------------------------------------------------

    if extension == "pdf":

        return extract_text_from_pdf(
            file_path
        )


    # --------------------------------------------------------
    # DOCX
    # --------------------------------------------------------

    if extension == "docx":

        return extract_text_from_docx(
            file_path
        )


    # --------------------------------------------------------
    # IMAGE
    # --------------------------------------------------------

    if extension in (
        SUPPORTED_IMAGE_EXTENSIONS
    ):

        return extract_text_from_image(
            file_path
        )


    # --------------------------------------------------------
    # UNSUPPORTED
    # --------------------------------------------------------

    print(
        f"Unsupported file type: "
        f"{extension}"
    )

    return ""


# ============================================================
# SIMPLE TEST
# ============================================================

if __name__ == "__main__":

    print(
        "\n========================================"
    )

    print(
        "DOCUMENT TEXT EXTRACTOR TEST"
    )

    print(
        "========================================"
    )


    # --------------------------------------------------------
    # CHECK TESSERACT
    # --------------------------------------------------------

    check_tesseract()


    # --------------------------------------------------------
    # FILE INPUT
    # --------------------------------------------------------

    test_file = input(

        "\nEnter full file path: "

    ).strip()


    # --------------------------------------------------------
    # EXTRACT TEXT
    # --------------------------------------------------------

    extracted_text = (
        extract_text_from_file(
            test_file
        )
    )


    # --------------------------------------------------------
    # RESULT
    # --------------------------------------------------------

    print(
        "\n========================================"
    )

    print(
        "EXTRACTED TEXT"
    )

    print(
        "========================================\n"
    )


    print(
        extracted_text
    )


    print(
        "\n========================================"
    )

    print(
        f"TOTAL CHARACTERS: "
        f"{len(extracted_text)}"
    )

    print(
        "========================================"
    )