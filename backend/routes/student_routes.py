# =========================================================
# routes/student_routes.py
# =========================================================

from flask import (
    Blueprint,
    request,
    jsonify,
    current_app
)

from config import db

from services.skill_extractor import (
    extract_skills_from_documents
)

from services.llm_worker import (
    start_background_llm_analysis
)

from werkzeug.security import (
    generate_password_hash,
    check_password_hash
)

from werkzeug.utils import secure_filename

from bson import ObjectId
from bson.errors import InvalidId

import jwt
import os
import uuid
import json

from datetime import (
    datetime,
    timedelta,
    timezone
)

from functools import wraps


# =========================================================
# OPTIONAL DOCUMENT EXTRACTION LIBRARIES
# =========================================================

try:

    from pypdf import PdfReader

    PDF_AVAILABLE = True

except ImportError:

    PDF_AVAILABLE = False


try:

    from docx import Document

    DOCX_AVAILABLE = True

except ImportError:

    DOCX_AVAILABLE = False


# =========================================================
# BLUEPRINT
# =========================================================

student_routes = Blueprint(
    "student_routes",
    __name__
)


# =========================================================
# MONGODB COLLECTION
# =========================================================

students_collection = db["students"]


# =========================================================
# JWT SECRET KEY
# =========================================================

JWT_SECRET_KEY = os.getenv(
    "JWT_SECRET_KEY"
)

if not JWT_SECRET_KEY:

    raise RuntimeError(
        "JWT_SECRET_KEY is not configured"
    )


# =========================================================
# ALLOWED FILE EXTENSIONS
# =========================================================

ALLOWED_RESUME_EXTENSIONS = {
    "pdf",
    "docx"
}

ALLOWED_CERTIFICATE_EXTENSIONS = {
    "pdf",
    "docx"
}


# =========================================================
# MAXIMUM FILE SIZE
# =========================================================

MAX_FILE_SIZE = 10 * 1024 * 1024
# 10 MB per file


# =========================================================
# HELPER - CURRENT UTC TIME
# =========================================================

def utc_now():

    return datetime.now(
        timezone.utc
    )


# =========================================================
# HELPER - FILE EXTENSION
# =========================================================

def get_file_extension(filename):

    if not filename:

        return ""

    filename = filename.lower()

    if "." not in filename:

        return ""

    return filename.rsplit(
        ".",
        1
    )[1]


# =========================================================
# HELPER - CHECK ALLOWED EXTENSION
# =========================================================

def allowed_file(
    filename,
    allowed_extensions
):

    extension = get_file_extension(
        filename
    )

    return (
        extension in
        allowed_extensions
    )


# =========================================================
# HELPER - EXTRACT PDF TEXT
# =========================================================

def extract_pdf_text(file_path):

    if not PDF_AVAILABLE:

        return ""

    try:

        reader = PdfReader(
            file_path
        )

        extracted_text = []

        for page in reader.pages:

            text = page.extract_text()

            if text:

                extracted_text.append(
                    text
                )

        return "\n".join(
            extracted_text
        ).strip()

    except Exception as error:

        print(
            "PDF extraction error:",
            error
        )

        return ""


# =========================================================
# HELPER - EXTRACT DOCX TEXT
# =========================================================

def extract_docx_text(file_path):

    if not DOCX_AVAILABLE:

        return ""

    try:

        document = Document(
            file_path
        )

        extracted_text = []

        # -------------------------------------------------
        # Paragraphs
        # -------------------------------------------------

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:

                extracted_text.append(
                    text
                )

        # -------------------------------------------------
        # Tables
        # -------------------------------------------------

        for table in document.tables:

            for row in table.rows:

                row_text = []

                for cell in row.cells:

                    text = cell.text.strip()

                    if text:

                        row_text.append(
                            text
                        )

                if row_text:

                    extracted_text.append(
                        " | ".join(
                            row_text
                        )
                    )

        return "\n".join(
            extracted_text
        ).strip()

    except Exception as error:

        print(
            "DOCX extraction error:",
            error
        )

        return ""


# =========================================================
# HELPER - EXTRACT DOCUMENT TEXT
# =========================================================

def extract_document_text(
    file_path,
    filename
):

    extension = get_file_extension(
        filename
    )

    if extension == "pdf":

        return extract_pdf_text(
            file_path
        )

    if extension == "docx":

        return extract_docx_text(
            file_path
        )

    return ""


# =========================================================
# HELPER - SAVE UPLOADED FILE
# =========================================================

def save_uploaded_file(
    uploaded_file,
    folder,
    allowed_extensions
):

    if not uploaded_file:

        return None

    original_filename = (
        uploaded_file.filename
    )

    if not original_filename:

        return None

    if not allowed_file(
        original_filename,
        allowed_extensions
    ):

        raise ValueError(
            "Unsupported file type. "
            "Only PDF and DOCX files are supported."
        )

    # -----------------------------------------------------
    # Secure filename
    # -----------------------------------------------------

    safe_filename = secure_filename(
        original_filename
    )

    if not safe_filename:

        raise ValueError(
            "Invalid filename"
        )

    # -----------------------------------------------------
    # Create unique filename
    # -----------------------------------------------------

    unique_filename = (
        uuid.uuid4().hex
        + "_"
        + safe_filename
    )

    file_path = os.path.join(
        folder,
        unique_filename
    )

    # -----------------------------------------------------
    # Save file
    # -----------------------------------------------------

    uploaded_file.save(
        file_path
    )

    # -----------------------------------------------------
    # Check file size
    # -----------------------------------------------------

    file_size = os.path.getsize(
        file_path
    )

    if file_size > MAX_FILE_SIZE:

        try:

            os.remove(
                file_path
            )

        except Exception:

            pass

        raise ValueError(
            "File size exceeds 10 MB"
        )

    # -----------------------------------------------------
    # Extract text
    # -----------------------------------------------------

    extracted_text = (
        extract_document_text(
            file_path,
            original_filename
        )
    )

    return {

        "original_filename":
            original_filename,

        "stored_filename":
            unique_filename,

        "file_path":
            file_path,

        "file_size":
            file_size,

        "file_type":
            get_file_extension(
                original_filename
            ),

        "extracted_text":
            extracted_text,

        "uploaded_at":
            utc_now()
    }


# =========================================================
# JWT TOKEN VERIFICATION
# =========================================================

def token_required(function):

    @wraps(function)
    def decorated(
        *args,
        **kwargs
    ):

        auth_header = request.headers.get(
            "Authorization"
        )

        if not auth_header:

            return jsonify({

                "success": False,

                "message":
                    "Authorization token is required"

            }), 401

        if not auth_header.startswith(
            "Bearer "
        ):

            return jsonify({

                "success": False,

                "message":
                    "Invalid authorization format"

            }), 401

        token = auth_header.split(
            " ",
            1
        )[1]

        try:

            payload = jwt.decode(
                token,
                JWT_SECRET_KEY,
                algorithms=["HS256"]
            )

            student_id = payload.get(
                "student_id"
            )

            if not student_id:

                return jsonify({

                    "success": False,

                    "message":
                        "Invalid token"

                }), 401

        except jwt.ExpiredSignatureError:

            return jsonify({

                "success": False,

                "message":
                    "Token has expired"

            }), 401

        except jwt.InvalidTokenError:

            return jsonify({

                "success": False,

                "message":
                    "Invalid token"

            }), 401

        return function(
            student_id,
            *args,
            **kwargs
        )

    return decorated


# =========================================================
# STUDENT REGISTRATION
# =========================================================

@student_routes.route(
    "/register",
    methods=["POST"]
)
def register_student():

    # =====================================================
    # GET STUDENT DATA
    # =====================================================

    data = None

    # -----------------------------------------------------
    # JSON
    # -----------------------------------------------------

    if request.is_json:

        data = request.get_json(
            silent=True
        )

    # -----------------------------------------------------
    # MULTIPART FORM
    # -----------------------------------------------------

    else:

        student_data = request.form.get(
            "student_data"
        )

        if student_data:

            try:

                data = json.loads(
                    student_data
                )

            except Exception:

                return jsonify({

                    "success": False,

                    "message":
                        "Invalid student_data JSON"

                }), 400

    if not data:

        return jsonify({

            "success": False,

            "message":
                "Invalid JSON data"

        }), 400


    # =====================================================
    # DEBUG
    # =====================================================

    print(
        "\n========================================"
    )

    print(
        "DATA RECEIVED FROM FRONTEND"
    )

    print(
        "========================================"
    )

    print(
        data
    )

    print(
        "========================================\n"
    )


    # =====================================================
    # REQUIRED FIELDS
    # =====================================================

    required_fields = [

        "full_name",

        "register_number",

        "email",

        "password"

    ]

    for field in required_fields:

        if not data.get(field):

            return jsonify({

                "success": False,

                "message":
                    f"{field} is required"

            }), 400


    # =====================================================
    # EXTRACT NESTED DATA
    # =====================================================

    academic_data = data.get(
        "academic",
        {}
    )

    if not isinstance(
        academic_data,
        dict
    ):

        academic_data = {}


    career_data = data.get(
        "career_preferences",
        {}
    )

    if not isinstance(
        career_data,
        dict
    ):

        career_data = {}


    resume_data = data.get(
        "resume",
        {}
    )

    if not isinstance(
        resume_data,
        dict
    ):

        resume_data = {}


    skills_data = data.get(
        "skills",
        []
    )

    if not isinstance(
        skills_data,
        list
    ):

        skills_data = []


    certifications_data = data.get(
        "certifications",
        []
    )

    if not isinstance(
        certifications_data,
        list
    ):

        certifications_data = []


    projects_data = data.get(
        "projects",
        []
    )

    if not isinstance(
        projects_data,
        list
    ):

        projects_data = []


    # =====================================================
    # CHECK EXISTING STUDENT
    # =====================================================

    existing_student = (
        students_collection.find_one({

            "$or": [

                {
                    "personal.email":
                        data["email"]
                },

                {
                    "personal.register_number":
                        data["register_number"]
                }

            ]

        })
    )


    if existing_student:

        return jsonify({

            "success": False,

            "message":
                "Student already registered"

        }), 409


    # =====================================================
    # UPLOAD DIRECTORIES
    # =====================================================

    resume_folder = current_app.config.get(
        "RESUME_FOLDER"
    )

    certificate_folder = current_app.config.get(
        "CERTIFICATE_FOLDER"
    )


    # =====================================================
    # UPLOAD RESUME
    # =====================================================

    uploaded_resume = None

    resume_file = request.files.get(
        "resume"
    )

    if resume_file:

        if not resume_folder:

            return jsonify({

                "success": False,

                "message":
                    "Resume upload folder is not configured"

            }), 500

        try:

            uploaded_resume = (
                save_uploaded_file(

                    resume_file,

                    resume_folder,

                    ALLOWED_RESUME_EXTENSIONS

                )
            )

        except ValueError as error:

            return jsonify({

                "success": False,

                "message":
                    str(error)

            }), 400


    # =====================================================
    # UPLOAD CERTIFICATES
    # =====================================================

    uploaded_certificates = []

    certificate_files = (
        request.files.getlist(
            "certificates"
        )
    )

    if certificate_files:

        if not certificate_folder:

            return jsonify({

                "success": False,

                "message":
                    "Certificate upload folder is not configured"

            }), 500

        for certificate_file in certificate_files:

            if not certificate_file.filename:

                continue

            try:

                certificate_info = (
                    save_uploaded_file(

                        certificate_file,

                        certificate_folder,

                        ALLOWED_CERTIFICATE_EXTENSIONS

                    )
                )

                if certificate_info:

                    uploaded_certificates.append(
                        certificate_info
                    )

            except ValueError as error:

                return jsonify({

                    "success": False,

                    "message":
                        str(error)

                }), 400


    # =====================================================
    # DOCUMENT TEXT
    # =====================================================

    resume_extracted_text = ""

    if uploaded_resume:

        resume_extracted_text = (
            uploaded_resume.get(
                "extracted_text",
                ""
            )
            or ""
        )


    certificate_extracted_texts = []

    for certificate in uploaded_certificates:

        extracted_text = (
            certificate.get(
                "extracted_text",
                ""
            )
            or ""
        )

        if extracted_text:

            certificate_extracted_texts.append(
                extracted_text
            )


    # =====================================================
    # AUTOMATIC SKILL EXTRACTION
    # =====================================================

    try:

        detected_skills = (
            extract_skills_from_documents(

                resume_text=
                    resume_extracted_text,

                certificate_texts=
                    certificate_extracted_texts

            )
        )

    except Exception as error:

        print(
            "Skill extraction error:",
            error
        )

        detected_skills = {

            "skills": [],

            "categorized_skills": {},

            "skill_details": [],

            "document_count": 0

        }


    # =====================================================
    # DEBUG SKILLS
    # =====================================================

    print(
        "\n========================================"
    )

    print(
        "AUTOMATICALLY DETECTED SKILLS"
    )

    print(
        "========================================"
    )

    print(
        detected_skills
    )

    print(
        "========================================\n"
    )


    # =====================================================
    # MERGE CERTIFICATE INFORMATION
    # =====================================================

    final_certifications = []

    for index, certification in enumerate(
        certifications_data
    ):

        if not isinstance(
            certification,
            dict
        ):

            continue

        certification_copy = dict(
            certification
        )

        if index < len(
            uploaded_certificates
        ):

            certification_copy[
                "file"
            ] = uploaded_certificates[
                index
            ]

        final_certifications.append(
            certification_copy
        )


    # =====================================================
    # RESUME DATA
    # =====================================================

    final_resume = {

        "has_resume":
            resume_data.get(
                "has_resume"
            ),

        "resume_name":
            resume_data.get(
                "resume_name"
            ),

        "file":
            uploaded_resume

    }


    # =====================================================
    # CREATE STUDENT DOCUMENT
    # =====================================================

    student = {

        # =================================================
        # PERSONAL
        # =================================================

        "personal": {

            "full_name":
                data.get(
                    "full_name"
                ),

            "register_number":
                data.get(
                    "register_number"
                ),

            "roll_number":
                data.get(
                    "roll_number"
                ),

            "email":
                data.get(
                    "email"
                ),

            "mobile":
                data.get(
                    "mobile"
                ),

            "gender":
                data.get(
                    "gender"
                ),

            "date_of_birth":
                data.get(
                    "date_of_birth"
                ),

            "department":
                data.get(
                    "department"
                ),

            "degree":
                data.get(
                    "degree"
                ),

            "year_of_study":
                data.get(
                    "year_of_study"
                ),

            "section":
                data.get(
                    "section"
                ),

            "semester":
                data.get(
                    "semester"
                )

        },


        # =================================================
        # ACADEMIC
        # =================================================

        "academic": {

            "college_name":
                academic_data.get(
                    "college_name"
                ),

            "university":
                academic_data.get(
                    "university"
                ),

            "branch":
                academic_data.get(
                    "branch"
                ),

            "current_cgpa":
                academic_data.get(
                    "current_cgpa"
                ),

            "tenth_percentage":
                academic_data.get(
                    "tenth_percentage"
                ),

            "twelfth_percentage":
                academic_data.get(
                    "twelfth_percentage"
                ),

            "number_of_arrears":
                academic_data.get(
                    "number_of_arrears"
                ),

            "backlog_history":
                academic_data.get(
                    "backlog_history"
                ),

            "academic_year":
                academic_data.get(
                    "academic_year"
                ),

            "graduation_year":
                academic_data.get(
                    "graduation_year"
                )

        },


        # =================================================
        # PREFERRED / SELF-DECLARED SKILLS
        # =================================================

        "skills":
            skills_data,


        # =================================================
        # AUTOMATICALLY DETECTED SKILLS
        # =================================================

        "detected_skills": {

            "resume": (
                detected_skills.get(
                    "resume_skills",
                    []
                )
                if isinstance(
                    detected_skills,
                    dict
                )
                else []
            ),

            "certificates": (
                detected_skills.get(
                    "certificate_skills",
                    []
                )
                if isinstance(
                    detected_skills,
                    dict
                )
                else []
            ),

            "all": (
                detected_skills.get(
                    "skills",
                    []
                )
                if isinstance(
                    detected_skills,
                    dict
                )
                else []
            ),

            "details": (
                detected_skills.get(
                    "skill_details",
                    []
                )
                if isinstance(
                    detected_skills,
                    dict
                )
                else []
            )

        },


        # =================================================
        # CERTIFICATIONS
        # =================================================

        "certifications":
            final_certifications,


        # =================================================
        # RESUME
        # =================================================

        "resume":
            final_resume,


        # =================================================
        # PROJECTS
        # =================================================

        "projects":
            projects_data,


        # =================================================
        # CAREER PREFERENCES
        # =================================================

        "career_preferences": {

            "interested_domain":
                career_data.get(
                    "interested_domain"
                ),

            "preferred_job_role":
                career_data.get(
                    "preferred_job_role"
                ),

            "preferred_location":
                career_data.get(
                    "preferred_location"
                ),

            "internship_preferences":
                career_data.get(
                    "internship_preferences"
                ),

            "career_goal":
                career_data.get(
                    "career_goal"
                ),

            "learning_goal":
                career_data.get(
                    "learning_goal"
                )

        },


        # =================================================
        # LLM ANALYSIS
        # =================================================

        "llm_analysis": {

            "status":
                "pending",

            "skill_summary":
                None,

            "overall_score":
                None,

            "strengths":
                [],

            "skill_gaps":
                [],

            "career_analysis":
                None,

            "learning_recommendations":
                [],

            "overall_assessment":
                None

        },


        # =================================================
        # JOB RECOMMENDATIONS
        # =================================================

        "job_recommendations":
            [],


        # =================================================
        # PASSWORD
        # =================================================

        "password":
            generate_password_hash(
                data["password"]
            ),


        # =================================================
        # CREATED DATE
        # =================================================

        "created_at":
            utc_now()

    }


    # =====================================================
    # SAVE STUDENT TO MONGODB
    #
    # IMPORTANT:
    # Student is saved BEFORE starting Ollama.
    # =====================================================

    try:

        result = (
            students_collection.insert_one(
                student
            )
        )

    except Exception as error:

        print(
            "MongoDB insertion error:",
            error
        )

        return jsonify({

            "success": False,

            "message":
                "Failed to save student"

        }), 500


    # =====================================================
    # STUDENT ID
    # =====================================================

    student_id = str(
        result.inserted_id
    )


    # =====================================================
    # START BACKGROUND LLM ANALYSIS
    #
    # IMPORTANT:
    # Do NOT call analyze_student() here.
    # =====================================================

    try:

        start_background_llm_analysis(
            student_id
        )

        print(
            "Background LLM analysis "
            "queued successfully."
        )

    except Exception as error:

        print(
            "Could not start background "
            "LLM analysis:",
            error
        )

        # -----------------------------------------------
        # Registration is already successful.
        # Save the worker error for debugging.
        # -----------------------------------------------

        students_collection.update_one(

            {
                "_id":
                    result.inserted_id
            },

            {
                "$set": {

                    "llm_analysis.status":
                        "failed",

                    "llm_analysis.error":
                        str(error)

                }
            }

        )


    # =====================================================
    # REGISTRATION SUCCESS
    #
    # THIS RESPONSE DOES NOT WAIT FOR OLLAMA.
    # =====================================================

    return jsonify({

        "success": True,

        "message":
            "Student registered successfully",

        "student_id":
            student_id,

        "llm_analysis_status":
            "pending",

        "resume_uploaded":
            uploaded_resume is not None,

        "certificates_uploaded":
            len(
                uploaded_certificates
            ),

        "resume_text_extracted":
            bool(

                uploaded_resume

                and uploaded_resume.get(
                    "extracted_text"
                )

            ),

        "certificate_texts_extracted":
            sum(

                1

                for certificate
                in uploaded_certificates

                if certificate.get(
                    "extracted_text"
                )

            ),

        "skills_detected":
            len(

                detected_skills.get(
                    "skills",
                    []
                )

                if isinstance(
                    detected_skills,
                    dict
                )

                else []

            )

    }), 201


# =========================================================
# STUDENT LOGIN
# =========================================================

@student_routes.route(
    "/login",
    methods=["POST"]
)
def login_student():

    data = request.get_json(
        silent=True
    )

    if not data:

        return jsonify({

            "success": False,

            "message":
                "Invalid JSON data"

        }), 400


    email = data.get(
        "email"
    )

    password = data.get(
        "password"
    )


    if not email or not password:

        return jsonify({

            "success": False,

            "message":
                "Email and password are required"

        }), 400


    student = (
        students_collection.find_one({

            "personal.email":
                email

        })
    )


    if not student:

        return jsonify({

            "success": False,

            "message":
                "Invalid email or password"

        }), 401


    stored_password = student.get(
        "password"
    )


    if not stored_password:

        return jsonify({

            "success": False,

            "message":
                "Password is not configured"

        }), 500


    try:

        password_valid = (
            check_password_hash(

                stored_password,

                password

            )
        )

    except Exception:

        password_valid = False


    if not password_valid:

        return jsonify({

            "success": False,

            "message":
                "Invalid email or password"

        }), 401


    payload = {

        "student_id":
            str(
                student["_id"]
            ),

        "email":
            student["personal"]["email"],

        "exp":
            utc_now()
            +
            timedelta(
                hours=24
            )

    }


    token = jwt.encode(

        payload,

        JWT_SECRET_KEY,

        algorithm="HS256"

    )


    return jsonify({

        "success": True,

        "message":
            "Login successful",

        "token":
            token,

        "student": {

            "student_id":
                str(
                    student["_id"]
                ),

            "full_name":
                student["personal"][
                    "full_name"
                ],

            "email":
                student["personal"][
                    "email"
                ],

            "register_number":
                student["personal"][
                    "register_number"
                ]

        }

    }), 200


# =========================================================
# GET LOGGED-IN STUDENT PROFILE
# =========================================================

@student_routes.route(
    "/profile",
    methods=["GET"]
)
@token_required
def get_student_profile(
    student_id
):

    try:

        object_id = ObjectId(
            student_id
        )

    except InvalidId:

        return jsonify({

            "success": False,

            "message":
                "Invalid student ID"

        }), 400


    student = (
        students_collection.find_one({

            "_id":
                object_id

        })
    )


    if not student:

        return jsonify({

            "success": False,

            "message":
                "Student not found"

        }), 404


    # -----------------------------------------------------
    # Never send password
    # -----------------------------------------------------

    student.pop(
        "password",
        None
    )


    student["_id"] = str(
        student["_id"]
    )


    return jsonify({

        "success": True,

        "student":
            student

    }), 200


# =========================================================
# GET LLM ANALYSIS STATUS / RESULT
# =========================================================

@student_routes.route(
    "/analysis",
    methods=["GET"]
)
@token_required
def get_student_analysis(
    student_id
):

    try:

        object_id = ObjectId(
            student_id
        )

    except InvalidId:

        return jsonify({

            "success": False,

            "message":
                "Invalid student ID"

        }), 400


    student = (
        students_collection.find_one(

            {
                "_id":
                    object_id
            },

            {
                "llm_analysis": 1
            }

        )
    )


    if not student:

        return jsonify({

            "success": False,

            "message":
                "Student not found"

        }), 404


    analysis = student.get(
        "llm_analysis",
        {}
    )


    return jsonify({

        "success": True,

        "analysis":
            analysis

    }), 200


# =========================================================
# UPDATE STUDENT PROFILE
# =========================================================

@student_routes.route(
    "/profile",
    methods=["PUT"]
)
@token_required
def update_student_profile(
    student_id
):

    data = request.get_json(
        silent=True
    )


    if not data:

        return jsonify({

            "success": False,

            "message":
                "Invalid JSON data"

        }), 400


    try:

        object_id = ObjectId(
            student_id
        )

    except InvalidId:

        return jsonify({

            "success": False,

            "message":
                "Invalid student ID"

        }), 400


    student = (
        students_collection.find_one({

            "_id":
                object_id

        })
    )


    if not student:

        return jsonify({

            "success": False,

            "message":
                "Student not found"

        }), 404


    # =====================================================
    # ALLOWED PROFILE FIELDS
    # =====================================================

    allowed_fields = [

        "personal",

        "academic",

        "skills",

        "detected_skills",

        "certifications",

        "resume",

        "projects",

        "career_preferences"

    ]


    update_data = {}


    for field in allowed_fields:

        if field in data:

            update_data[field] = data[
                field
            ]


    if not update_data:

        return jsonify({

            "success": False,

            "message":
                "No valid profile data provided"

        }), 400


    update_data[
        "updated_at"
    ] = utc_now()


    students_collection.update_one(

        {
            "_id":
                object_id
        },

        {
            "$set":
                update_data
        }

    )


    return jsonify({

        "success": True,

        "message":
            "Student profile updated successfully"

    }), 200